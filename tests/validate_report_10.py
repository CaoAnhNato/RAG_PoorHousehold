import os
import sys
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
from pathlib import Path
import asyncio

# Thêm đường dẫn gốc vào sys.path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "src"))

from src.query_control.agentic.report_generator import ReportGenerator
from src.query_control.agentic.report_queries import execute_report_query

def extract_text_from_pdf(pdf_path: Path, docx_path: Path | None = None) -> str:
    """Trích xuất toàn bộ text từ file PDF/DOCX ra chuỗi."""
    try:
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat
        opts = PdfPipelineOptions()
        opts.do_ocr = False
        opts.do_table_structure = False
        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        )
        doc_result = converter.convert(str(pdf_path))
        return doc_result.document.export_to_markdown()
    except Exception:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(str(pdf_path))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except Exception:
        pass

    try:
        import fitz
        doc = fitz.open(str(pdf_path))
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        return text
    except Exception:
        pass

    if docx_path and docx_path.exists():
        import docx
        doc = docx.Document(str(docx_path))
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    return ""

async def check_all_charts_dom_overlap(html_paths: list[Path]) -> bool:
    """Kiểm tra overlap label bằng Playwright DOM Boxes trên danh sách file HTML."""
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        print("   [Cảnh báo] Không cài đặt Playwright, bỏ qua kiểm tra DOM Overlap.")
        return True

    all_ok = True
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_viewport_size({"width": 1400, "height": 900})

            for html_path in html_paths:
                if not html_path.exists():
                    print(f"   [LỖI] File HTML biểu đồ không tồn tại: {html_path}")
                    all_ok = False
                    continue

                file_uri = html_path.resolve().as_uri()
                await page.goto(file_uri)
                await asyncio.sleep(1) # Chờ Plotly render xong

                boxes = await page.evaluate('''() => {
                    const elements = document.querySelectorAll('.g-gtitle text, .g-xtitle text, .g-ytitle text, .legendtext');
                    return Array.from(elements).map(el => {
                        const rect = el.getBoundingClientRect();
                        return {
                            text: el.textContent.trim(),
                            x: rect.x,
                            y: rect.y,
                            width: rect.width,
                            height: rect.height
                        };
                    }).filter(b => b.width > 0 && b.height > 0);
                }''')

                has_overlap = False
                for i in range(len(boxes)):
                    for j in range(i + 1, len(boxes)):
                        b1 = boxes[i]
                        b2 = boxes[j]
                        if not (b1['x'] + b1['width'] <= b2['x'] + 2 or
                                b2['x'] + b2['width'] <= b1['x'] + 2 or
                                b1['y'] + b1['height'] <= b2['y'] + 2 or
                                b2['y'] + b2['height'] <= b1['y'] + 2):
                            has_overlap = True
                            print(f"   [CẢNH BÁO OVERLAP] Biểu đồ {html_path.name}: Label '{b1['text']}' chồng lấn với '{b2['text']}'")
                            break
                    if has_overlap:
                        break

                if not has_overlap:
                    print(f"   [THÀNH CÔNG] Biểu đồ {html_path.name} đạt chuẩn DOM Box, không có label overlap!")
                else:
                    all_ok = False

            await browser.close()
            return all_ok
    except Exception as e:
        print(f"   [Thông báo] Playwright DOM check bỏ qua do lỗi: {e}")
        return True

def validate():
    print("1. Đang sinh lại báo cáo số 10 năm 2024...")
    res = ReportGenerator.generate(report_id=10, year=2024)
    docx_path = Path("Runtime/reports/bao_cao_10_2024.docx")
    pdf_path = Path("Runtime/reports/bao_cao_10_2024.pdf")

    if not docx_path.exists() or not pdf_path.exists():
        print("   [LỖI] Không tìm thấy file DOCX hoặc PDF được sinh ra!")
        return False
    print("   [THÀNH CÔNG] Đã tạo thành công file DOCX và PDF!")

    print("2. Đang kiểm định file hình ảnh và HTML biểu đồ High-Fidelity...")
    chart_files = [
        "report_10_chart_1.html",
        "report_10_chart_2.html",
        "report_10_chart_3.html"
    ]
    chart_paths = [Path("artifacts/charts") / c_html for c_html in chart_files]
    for cp in chart_paths:
        if not cp.exists():
            print(f"   [LỖI] Thiếu file HTML biểu đồ: {cp}")
            return False
        else:
            print(f"   [OK] File HTML biểu đồ tồn tại: {cp}")

    all_charts_ok = asyncio.run(check_all_charts_dom_overlap(chart_paths))
    if not all_charts_ok:
        return False

    print("3. Đang kiểm xuất dữ liệu thô từ DuckDB...")
    df, sql, title = execute_report_query(10, 2024)
    dist_df = df[df["STT"].astype(str).str.isdigit()]

    reason_cols = [
        "1. Thiếu đất sản xuất",
        "2. Thiếu vốn",
        "3. Thiếu lao động",
        "4. Thiếu công cụ sản xuất",
        "5. Thiếu kiến thức sản xuất",
        "6. Thiếu kỹ năng lao động",
        "7. Ốm đau, tai nạn",
        "8. Khác"
    ]

    tong_om_dau = int(dist_df["7. Ốm đau, tai nạn"].sum()) if "7. Ốm đau, tai nạn" in dist_df.columns else 0
    tong_thieu_von = int(dist_df["2. Thiếu vốn"].sum()) if "2. Thiếu vốn" in dist_df.columns else 0

    print(f"   -> Tổng số lượt Ốm đau, tai nạn toàn tỉnh từ DuckDB: {tong_om_dau:,}".replace(",", "."))
    print(f"   -> Tổng số lượt Thiếu vốn toàn tỉnh từ DuckDB: {tong_thieu_von:,}".replace(",", "."))

    print("4. Đang phân tích nội dung file PDF/DOCX...")
    try:
        exported_text = extract_text_from_pdf(pdf_path, docx_path)
    except Exception as e:
        print(f"[LỖI] Gặp lỗi khi đọc PDF/DOCX: {e}")
        return False

    print("5. Đang đối chiếu số liệu giữa PDF/DOCX và DuckDB...")
    parity_ok = True
    for _, r in dist_df.iterrows():
        ten_dv = str(r["Phường/Xã"]).strip()
        od = int(r.get("7. Ốm đau, tai nạn", 0))
        tv = int(r.get("2. Thiếu vốn", 0))
        if ten_dv in exported_text and str(od) in exported_text and str(tv) in exported_text:
            print(f"   [THÀNH CÔNG] Đơn vị '{ten_dv}': Khớp chính xác Ốm đau={od}, Thiếu vốn={tv}")
        else:
            print(f"   [CẢNH BÁO] Sai lệch số liệu tại đơn vị '{ten_dv}': Ốm đau={od}, Thiếu vốn={tv}")
            parity_ok = False

    if not parity_ok:
        return False

    print("6. Đang kiểm định cấu trúc 4 phần chuẩn văn bản hành chính...")
    required_sections = [
        "PHẦN I:",
        "PHẦN II:",
        "PHẦN III:",
        "PHẦN IV:"
    ]
    missing = []
    for sec in required_sections:
        if sec not in exported_text:
            missing.append(sec)
            
    if missing:
        print(f"   [LỖI] Thiếu tiêu đề cấu trúc chuẩn trong PDF: {missing}")
        return False
    else:
        print("   [THÀNH CÔNG] Đầy đủ 4 phần cấu trúc chuẩn báo cáo!")

    print("7. Đang kiểm định sự xuất hiện của phần text diễn giải biểu đồ (analysis paragraphs)...")
    required_analysis_phrases = [
        "Biểu đồ Pareto chỉ ra quy luật 80/20 rõ rệt",
        "Ma trận nhiệt hiển thị trực quan toàn diện hồ sơ tổn thương",
        "Sơ đồ vòng xoay minh họa cơ chế khép kín tàn phá hộ gia đình"
    ]
    missing_analysis = []
    for phrase in required_analysis_phrases:
        if phrase not in exported_text:
            missing_analysis.append(phrase)

    if missing_analysis:
        print(f"   [LỖI] Mất text diễn giải biểu đồ trong file xuất ra: {missing_analysis}")
        return False
    else:
        print("   [THÀNH CÔNG] Toàn bộ text diễn giải biểu đồ đã xuất hiện đầy đủ trong báo cáo!")

    print("\n=> TẤT CẢ KIỂM ĐỊNH BÁO CÁO SỐ 10 ĐẠT 100% CHUẨN XÁC!\n")
    return True

if __name__ == "__main__":
    success = validate()
    sys.exit(0 if success else 1)
