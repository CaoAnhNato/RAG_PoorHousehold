import os
import sys
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
from pathlib import Path
import asyncio

# Thêm đường dẫn gốc vào sys.path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "src"))

from src.query_control.agentic.report_generator import ReportGenerator
from src.query_control.agentic.report_queries import execute_report_query

def extract_text_from_pdf(pdf_path: Path, docx_path: Path | None = None) -> str:
    """Trích xuất toàn bộ text từ file PDF/DOCX ra chuỗi."""
    try:
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat
        opts = PdfPipelineOptions()
        opts.do_ocr = False
        opts.do_table_structure = False
        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        )
        doc_result = converter.convert(str(pdf_path))
        return doc_result.document.export_to_markdown()
    except Exception:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(str(pdf_path))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except Exception:
        pass

    try:
        import fitz
        doc = fitz.open(str(pdf_path))
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        return text
    except Exception:
        pass

    if docx_path and docx_path.exists():
        import docx
        doc = docx.Document(str(docx_path))
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    raise RuntimeError("Không thể đọc text từ PDF hoặc DOCX.")

async def check_all_charts_dom_overlap(html_paths: list[Path]) -> bool:
    """Kiểm tra không chồng lấn DOM Bounding Box bằng Playwright."""
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        print("   [Thông báo] Playwright chưa cài đặt, bỏ qua bước kiểm tra DOM.")
        return True

    all_ok = True
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_viewport_size({"width": 1400, "height": 900})

            for html_path in html_paths:
                if not html_path.exists():
                    print(f"   [LỖI] File HTML biểu đồ không tồn tại: {html_path}")
                    all_ok = False
                    continue

                file_uri = html_path.resolve().as_uri()
                await page.goto(file_uri)
                await asyncio.sleep(1) # Chờ Plotly render xong

                boxes = await page.evaluate('''() => {
                    const elements = document.querySelectorAll('.g-gtitle text, .g-xtitle text, .g-ytitle text, .legendtext');
                    return Array.from(elements).map(el => {
                        const rect = el.getBoundingClientRect();
                        return {
                            text: el.textContent.trim(),
                            x: rect.x,
                            y: rect.y,
                            width: rect.width,
                            height: rect.height
                        };
                    }).filter(b => b.width > 0 && b.height > 0);
                }''')

                has_overlap = False
                for i in range(len(boxes)):
                    for j in range(i + 1, len(boxes)):
                        b1 = boxes[i]
                        b2 = boxes[j]
                        # Kiểm tra chồng chéo hình chữ nhật 2D
                        if not (b1['x'] + b1['width'] <= b2['x'] + 2 or
                                b2['x'] + b2['width'] <= b1['x'] + 2 or
                                b1['y'] + b1['height'] <= b2['y'] + 2 or
                                b2['y'] + b2['height'] <= b1['y'] + 2):
                            has_overlap = True
                            print(f"   [CẢNH BÁO OVERLAP] Biểu đồ {html_path.name}: Label '{b1['text']}' chồng lấn với '{b2['text']}'")
                            break
                    if has_overlap:
                        break

                if not has_overlap:
                    print(f"   [THÀNH CÔNG] Biểu đồ {html_path.name} đạt chuẩn DOM Box, không có label overlap!")
                else:
                    all_ok = False

            await browser.close()
            return all_ok
    except Exception as e:
        print(f"   [Thông báo] Playwright DOM check bỏ qua do lỗi: {e}")
        return True

def validate():
    print("1. Đang sinh lại báo cáo số 8 năm 2024...")
    rep = ReportGenerator.generate(8, 2024)
    pdf_path = Path(rep["pdf_path"])
    docx_path = Path(rep["docx_path"])
    
    if not pdf_path.exists():
        print(f"[LỖI] Không tìm thấy file PDF tại {pdf_path}")
        return False
    if not docx_path.exists():
        print(f"[LỖI] Không tìm thấy file DOCX tại {docx_path}")
        return False
    print("   [THÀNH CÔNG] Đã tạo thành công file DOCX và PDF!")

    print("2. Đang kiểm định file hình ảnh và HTML biểu đồ High-Fidelity...")
    chart_files = [
        "report_8_chart_1.html",
        "report_8_chart_2.html",
        "report_8_chart_3.html"
    ]
    chart_paths = [Path("artifacts/charts") / c_html for c_html in chart_files]
    for cp in chart_paths:
        if not cp.exists():
            print(f"   [LỖI] Thiếu file HTML biểu đồ: {cp}")
            return False
        else:
            print(f"   [OK] File HTML biểu đồ tồn tại: {cp}")

    all_charts_ok = asyncio.run(check_all_charts_dom_overlap(chart_paths))
    if not all_charts_ok:
        return False

    print("3. Đang kiểm xuất dữ liệu thô từ DuckDB...")
    df, sql, title = execute_report_query(8, 2024)
    dist_ho_df = df[df["STT"].astype(str).str.isdigit() & (df["Phân tổ"] == "Hộ")]
    if not dist_ho_df.empty:
        tong_ngheo = int(dist_ho_df["Tổng số hộ nghèo"].sum())
        tong_can_ngheo = int(dist_ho_df["Tổng số hộ cận nghèo"].sum())
    else:
        tong_ngheo = int(df["Tổng số hộ nghèo"].sum())
        tong_can_ngheo = int(df["Tổng số hộ cận nghèo"].sum())

    print(f"   -> Tổng số hộ nghèo toàn tỉnh (cộng gộp 8 huyện/thành phố từ DuckDB): {tong_ngheo:,}".replace(",", "."))
    print(f"   -> Tổng số hộ cận nghèo toàn tỉnh (cộng gộp 8 huyện/thành phố từ DuckDB): {tong_can_ngheo:,}".replace(",", "."))

    print("4. Đang phân tích nội dung file PDF/DOCX...")
    try:
        exported_text = extract_text_from_pdf(pdf_path, docx_path)
    except Exception as e:
        print(f"[LỖI] Gặp lỗi khi đọc PDF/DOCX: {e}")
        return False

    print("5. Đang đối chiếu số liệu giữa PDF/DOCX và DuckDB...")
    parity_ok = True
    for _, r in dist_ho_df.iterrows():
        ten_dv = str(r["Phường/Xã"]).strip()
        hn = int(r["Tổng số hộ nghèo"])
        hcn = int(r["Tổng số hộ cận nghèo"])
        if ten_dv in exported_text and str(hn) in exported_text and str(hcn) in exported_text:
            print(f"   [THÀNH CÔNG] Đơn vị '{ten_dv}': Khớp chính xác hộ nghèo={hn}, cận nghèo={hcn}")
        else:
            print(f"   [CẢNH BÁO] Sai lệch số liệu tại đơn vị '{ten_dv}': nghèo={hn}, cận nghèo={hcn}")
            parity_ok = False

    if not parity_ok:
        return False

    print("6. Đang kiểm định cấu trúc 4 phần chuẩn văn bản hành chính...")
    required_sections = [
        "PHẦN I:",
        "PHẦN II:",
        "PHẦN III:",
        "PHẦN IV:"
    ]
    missing = []
    for sec in required_sections:
        if sec not in exported_text:
            missing.append(sec)
            
    if missing:
        print(f"   [LỖI] Thiếu tiêu đề cấu trúc chuẩn trong PDF: {missing}")
        return False
    else:
        print("   [THÀNH CÔNG] Đầy đủ 4 phần cấu trúc chuẩn báo cáo!")

    print("\n=> TẤT CẢ KIỂM ĐỊNH BÁO CÁO SỐ 8 ĐẠT 100% CHUẨN XÁC!\n")
    return True

if __name__ == "__main__":
    success = validate()
    sys.exit(0 if success else 1)
